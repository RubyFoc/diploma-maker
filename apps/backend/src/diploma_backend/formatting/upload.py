"""Formatting-sample upload storage + `.docx` parsing (TASK-E05-2).

Saves the raw uploaded `.docx` under `UPLOADS_DIR` and extracts a best-effort `InstitutionConfig`
page/font/citation-style seed from it. Does not touch MongoDB or FastAPI — `formatting.router`
wires this into the HTTP layer and `create_institution_config` (TASK-E05-1) persists the result.
"""

import os
import re
import uuid
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from diploma_backend.formatting.models import (
    CitationStyle,
    FontConfig,
    MarginsMm,
    Orientation,
    PageConfig,
    PageSize,
)

# Counting parent directories to find "the repo root" breaks across environments (local dev vs.
# the Docker image, where apps/backend/src is copied to /app/src — a different depth). Use an
# env var instead, defaulting to a path relative to the current working directory (repo root in
# local dev via `uv run` from apps/backend/, or /app in the container); `uploads/` is gitignored
# per AGENTS.md secrets policy regardless of where it resolves to.
UPLOADS_DIR = Path(os.environ.get("UPLOADS_DIR", "uploads"))

# Standard page sizes in mm, portrait orientation. Used to snap a scanned page size to the
# ADR-0005 `PageSize` enum ("A4 | Letter" only) by nearest match.
_A4_MM = (210.0, 297.0)
_LETTER_MM = (215.9, 279.4)

# Best-effort heuristics only (see `guess_citation_style` docstring): a numeric bracket like
# "[1]" is common to numeric reference-list styles (GOST-style academic writing in this
# platform's target institutions); "(Author, 2020)" is the APA in-text pattern. Neither regex
# is a real citation-style classifier.
_NUMERIC_CITATION_RE = re.compile(r"\[\d+\]")
_AUTHOR_YEAR_RE = re.compile(r"\([A-Za-zА-Яа-яЁё .,'-]+,\s*\d{4}\)")


class FormattingSampleParseError(ValueError):
    """Raised when an uploaded formatting sample cannot be parsed.

    Callers (the upload router) must translate this into a 4xx response rather than letting it
    surface as a 500 — per the epic's fail-closed policy, an unparseable sample must never be
    guessed at silently.
    """


@dataclass
class ParsedFormattingSample:
    """The subset of ADR-0005's `InstitutionConfig` fields extracted from a `.docx` sample."""

    page: PageConfig
    font: FontConfig
    citation_style: CitationStyle


def save_uploaded_sample(content: bytes) -> str:
    """Write `content` to `UPLOADS_DIR` under a generated id and return that id.

    The id (a UUID4 hex string) doubles as the file's basename (`<id>.docx`) and is what the
    caller stores as `raw_sample_reference`. Side effect: creates `UPLOADS_DIR` if missing and
    writes one file into it.
    """
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    file_id = uuid.uuid4().hex
    (UPLOADS_DIR / f"{file_id}.docx").write_bytes(content)
    return file_id


def parse_formatting_sample(content: bytes) -> ParsedFormattingSample:
    """Parse `content` (raw `.docx` bytes) into page/font/citation-style fields.

    Reads `document.sections[0]` for page size/orientation/margins, the `Normal` style (falling
    back to the first body paragraph's run) for the default font, and scans paragraph text for a
    citation-style guess (see `guess_citation_style`).

    Raises `FormattingSampleParseError` if `content` is not a valid `.docx` file, or if margins
    or a font family/size cannot be determined — this module fails closed rather than inventing
    plausible-looking defaults for data it couldn't read.
    """
    try:
        document = Document(BytesIO(content))
    except (PackageNotFoundError, KeyError, ValueError, zipfile.BadZipFile) as exc:
        raise FormattingSampleParseError("Uploaded file is not a valid .docx document") from exc

    page = _parse_page(document)
    font = _parse_font(document)
    citation_style = guess_citation_style(document)
    return ParsedFormattingSample(page=page, font=font, citation_style=citation_style)


def _parse_page(document: Document) -> PageConfig:
    section = document.sections[0]

    margins = {
        "top": section.top_margin,
        "bottom": section.bottom_margin,
        "left": section.left_margin,
        "right": section.right_margin,
    }
    if any(value is None for value in margins.values()):
        raise FormattingSampleParseError("Could not read page margins from the .docx sample")
    # python-docx stores lengths as EMUs; converting back to mm introduces sub-mm rounding noise
    # (e.g. 30.00375 for an exact 30mm margin) that isn't meaningful for formatting rules
    # specified in whole millimeters, so round to the nearest mm here.
    margins_mm = MarginsMm(**{key: round(value.mm) for key, value in margins.items()})

    if section.page_width is None or section.page_height is None:
        raise FormattingSampleParseError("Could not read page dimensions from the .docx sample")
    width_mm, height_mm = section.page_width.mm, section.page_height.mm
    orientation: Orientation = "landscape" if width_mm > height_mm else "portrait"
    size = _guess_page_size(width_mm, height_mm)

    return PageConfig(size=size, orientation=orientation, margins_mm=margins_mm)


def _guess_page_size(width_mm: float, height_mm: float) -> PageSize:
    """Snap `(width_mm, height_mm)` to the closer of A4/Letter (portrait-normalized)."""
    portrait_dims = (min(width_mm, height_mm), max(width_mm, height_mm))
    a4_diff = abs(portrait_dims[0] - _A4_MM[0]) + abs(portrait_dims[1] - _A4_MM[1])
    letter_diff = abs(portrait_dims[0] - _LETTER_MM[0]) + abs(portrait_dims[1] - _LETTER_MM[1])
    return "A4" if a4_diff <= letter_diff else "Letter"


def _parse_font(document: Document) -> FontConfig:
    normal_style = document.styles["Normal"]
    family = normal_style.font.name
    size = normal_style.font.size
    line_spacing = normal_style.paragraph_format.line_spacing

    if family is None or size is None:
        first_run = _first_body_run(document)
        if first_run is not None:
            family = family or first_run.font.name
            size = size or first_run.font.size

    if family is None or size is None:
        raise FormattingSampleParseError(
            "Could not determine a default font family/size from the .docx sample"
        )

    # Single spacing has no explicit `line_spacing` value in python-docx; treating that absence
    # as 1.0 is a reasonable default, not a guess about data we failed to read.
    return FontConfig(family=family, size_pt=size.pt, line_spacing=line_spacing or 1.0)


def _first_body_run(document: Document):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.name is not None or run.font.size is not None:
                return run
    return None


def guess_citation_style(document: Document) -> CitationStyle:
    """Best-effort citation-style guess from paragraph text — not a real classifier.

    Counts bracketed-numeric citations (`[1]`, common to numeric/GOST-style reference lists)
    against author-year parentheticals (`(Author, 2020)`, the APA in-text pattern) across the
    document's body paragraphs and returns whichever pattern occurs more, or `"custom"` if
    neither appears. This is a simple regex heuristic for the MVP, explicitly not an NLP
    citation-style classifier; it will misclassify styles it has no rule for (e.g. MLA) and any
    document that mixes conventions. Revisit only if real usage data shows this matters.
    """
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    numeric_hits = len(_NUMERIC_CITATION_RE.findall(text))
    author_year_hits = len(_AUTHOR_YEAR_RE.findall(text))

    if numeric_hits == 0 and author_year_hits == 0:
        return "custom"
    return "GOST" if numeric_hits >= author_year_hits else "APA"
