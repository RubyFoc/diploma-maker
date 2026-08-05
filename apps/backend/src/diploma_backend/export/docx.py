"""Markdown → `.docx` mapping engine (TASK-E06-1) + institution config styling (TASK-E06-2).

Converts LLM-generated thesis-chapter Markdown into a `docx.Document`. `TASK-E06-3` (media
placeholders) builds on this later; that extension is not implemented here.

Supported Markdown subset (the platform's LLM content generation is expected to only ever
produce this subset, not arbitrary web Markdown):
- Headings `#`, `##`, `###` → Word's built-in "Heading 1"/"Heading 2"/"Heading 3" paragraph
  styles.
- Blank-line-separated blocks of text → plain paragraphs.
- Inline `**bold**` and `*italic*` within any paragraph/heading/list-item text → `run.bold` /
  `run.italic` on the corresponding runs (formatting is preserved, not stripped).
- Unordered lists (`- item` / `* item`) → Word's built-in "List Bullet" style, one paragraph per
  item.
- Ordered lists (`1. item`, `2. item`, ...) → Word's built-in "List Number" style, one paragraph
  per item. The actual number text is not read; Word's own list numbering is used.

Explicitly NOT supported: tables, images/media (media placeholders are TASK-E06-3), nested
lists, links, blockquotes, code blocks, and headings beyond level 3 (`####`+). Any line matching
none of the rules above (including all of the above) is not dropped: it is emitted as a plain
paragraph containing the literal source text (markers and all), so worst case a document renders
literal Markdown syntax visibly rather than silently losing content.

`apply_institution_config` (TASK-E06-2) applies an `InstitutionConfig` (ADR-0005) to an already
built `Document`, in place:
- `page.size`/`page.orientation`/`page.margins_mm` → `document.sections[0]` width/height/margins.
- `font.family`/`font.size_pt`/`font.line_spacing` → the "Normal" style, so all body paragraphs
  and list items inherit it via Word's own style inheritance (not set per-paragraph).
- `headings.h1/h2/h3` → the "Heading 1"/"Heading 2"/"Heading 3" styles' font. Each `HeadingStyle`
  is an open-ended dict (ADR-0005 defines `{}`, no fixed keys); for the MVP only two keys are
  interpreted — `font_size_pt` (points) and `bold` (bool) — applied to the style's font. Any
  other key present in a `HeadingStyle` is ignored, not an error.
"""

import re
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Mm, Pt
from docx.text.paragraph import Paragraph

from diploma_backend.formatting.models import HeadingStyle, InstitutionConfig

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_UNORDERED_LIST_RE = re.compile(r"^[-*]\s+(.*)$")
_ORDERED_LIST_RE = re.compile(r"^\d+\.\s+(.*)$")
_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")

# Standard page sizes in mm, portrait orientation — matching `formatting/upload.py`'s
# `_A4_MM`/`_LETTER_MM` so parsing an uploaded sample and applying a config round-trip to the
# same dimensions.
_A4_MM = (210.0, 297.0)
_LETTER_MM = (215.9, 279.4)


def markdown_to_docx(markdown_text: str) -> Document:
    """Convert `markdown_text` into a `docx.Document` per the module docstring's supported subset.

    Parses `markdown_text` line by line (no Markdown AST library — see module docstring for the
    exact rules) and builds the document by appending one paragraph per heading, list item, or
    blank-line-separated paragraph block. Never raises for unsupported Markdown constructs: they
    fall through to being rendered as a plain paragraph containing the literal source text.
    """
    document = Document()
    paragraph_buffer: list[str] = []

    def flush_paragraph_buffer() -> None:
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, " ".join(paragraph_buffer))
        paragraph_buffer.clear()

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()

        if line == "":
            flush_paragraph_buffer()
            continue

        heading_match = _HEADING_RE.match(line)
        if heading_match:
            flush_paragraph_buffer()
            level = len(heading_match.group(1))
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline_runs(paragraph, heading_match.group(2))
            continue

        unordered_match = _UNORDERED_LIST_RE.match(line)
        if unordered_match:
            flush_paragraph_buffer()
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, unordered_match.group(1))
            continue

        ordered_match = _ORDERED_LIST_RE.match(line)
        if ordered_match:
            flush_paragraph_buffer()
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_runs(paragraph, ordered_match.group(1))
            continue

        # Anything else (plain text, and unsupported constructs like tables/blockquotes/code
        # fences) accumulates as part of the current paragraph block, per the module docstring's
        # fail-safe rule of rendering unsupported input as literal plain-paragraph text.
        paragraph_buffer.append(line)

    flush_paragraph_buffer()
    return document


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Convert `markdown_text` to `.docx` bytes, via `markdown_to_docx` and an in-memory buffer."""
    document = markdown_to_docx(markdown_text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def apply_institution_config(document: Document, config: InstitutionConfig) -> None:
    """Apply `config`'s page/font/heading styles to `document`, mutating it in place.

    Sets `document.sections[0]`'s page size/orientation/margins, the "Normal" style's font
    (family/size/line spacing, inherited by every body paragraph and list item), and the
    "Heading 1"/"Heading 2"/"Heading 3" styles' font. See the module docstring for exactly which
    `HeadingStyle` keys are interpreted (`font_size_pt`, `bold`) — anything else in those
    open-ended dicts is silently ignored, not an error.
    """
    _apply_page_config(document, config)
    _apply_font_config(document, config)
    _apply_heading_style(document, "Heading 1", config.headings.h1)
    _apply_heading_style(document, "Heading 2", config.headings.h2)
    _apply_heading_style(document, "Heading 3", config.headings.h3)


def _apply_page_config(document: Document, config: InstitutionConfig) -> None:
    section = document.sections[0]
    page = config.page

    width_mm, height_mm = _A4_MM if page.size == "A4" else _LETTER_MM
    if page.orientation == "landscape":
        width_mm, height_mm = height_mm, width_mm
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(width_mm)
    section.page_height = Mm(height_mm)

    margins = page.margins_mm
    section.top_margin = Mm(margins.top)
    section.bottom_margin = Mm(margins.bottom)
    section.left_margin = Mm(margins.left)
    section.right_margin = Mm(margins.right)


def _apply_font_config(document: Document, config: InstitutionConfig) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = config.font.family
    normal_style.font.size = Pt(config.font.size_pt)
    normal_style.paragraph_format.line_spacing = config.font.line_spacing


def _apply_heading_style(document: Document, style_name: str, heading_style: HeadingStyle) -> None:
    style_font = document.styles[style_name].font
    extras = heading_style.model_dump()

    font_size_pt = extras.get("font_size_pt")
    if font_size_pt is not None:
        style_font.size = Pt(font_size_pt)

    bold = extras.get("bold")
    if bold is not None:
        style_font.bold = bold


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    """Append `text` to `paragraph` as one or more runs, applying `**bold**`/`*italic*` markup.

    Splits `text` on the first of either marker (non-overlapping, left to right) and adds a run
    per segment, setting `run.bold`/`run.italic` only for matched segments — unmarked segments
    get a plain run with no formatting flags touched.
    """
    position = 0
    for match in _INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])

        bold_text, italic_text = match.group(1), match.group(2)
        if bold_text is not None:
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            run = paragraph.add_run(italic_text)
            run.italic = True

        position = match.end()

    if position < len(text):
        paragraph.add_run(text[position:])
