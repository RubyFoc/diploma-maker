"""`.docx` table-of-contents parsing (TASK-E10-2).

Extracts an ordered list of chapter titles from an uploaded `.docx` file. A thesis TOC upload on
this platform is really "tell me your chapter structure," not necessarily Word's auto-generated
TOC field, so this module supports six shapes of input, tried in order, each requiring
progressively less structure in the source file:

1. A sequence of `Heading 1`-styled paragraphs (mirroring how a real thesis document's actual
   chapter headings look) — preferred when present, since it's the least ambiguous signal.
2. Top-level entries from Word's auto-generated Table of Contents *field* — when a user builds
   their TOC via Word's References -> Table of Contents feature (common for a Cyrillic
   "оглавление" page extracted on its own, separate from the full thesis), each entry paragraph
   carries a `"TOC 1"`/`"TOC1"`-style name (level 2+ entries get `"TOC 2"` etc., which this
   deliberately ignores — same "only top-level" posture as the `Heading 1` case), not `Heading 1`.
3. Top-level paragraphs using Word's *list numbering* (a numbered-list style/button, not typed
   digits) — the rendered "1.", "2.", ... never appears in `paragraph.text` at all, since Word
   generates it from the list definition (`w:numPr`) rather than storing it as text, so neither
   of the text-based checks above nor the numbered-line regex below can see it. Detected via the
   paragraph's raw `w:numPr`/`w:ilvl` XML instead (level 0 only, i.e. not an indented sub-item).
4. Plain numbered lines like `"1. Introduction"`, `"2) Literature Review"`, or `"3 Conclusion"`,
   combined with (5) unnumbered lines that have a page-number suffix instead (a dot-leader or tab
   before the trailing digits) — e.g. a thesis's unnumbered front/back matter ("ВВЕДЕНИЕ",
   "ЗАКЛЮЧЕНИЕ", "СПИСОК ЛИТЕРАТУРЫ") sitting alongside numbered chapters, scanned together in
   one pass so neither shape drops the other.
6. Absolute last resort: every remaining non-blank line, when a real TOC was typed by hand with
   no styling, numbering, or page numbers at all — the only signal left at that point is that the
   caller chose this specific "upload a table of contents" endpoint, which this module trusts
   (excluding lines recognizable as a subsection or the page's own heading — see
   `_is_toc_subsection_or_page_title`).

Raises `TocParseError` only if `content` isn't a valid `.docx` file, or the document has no
non-blank paragraphs at all — tier 6 means an uploaded document essentially always yields
*something*, by design (see that tier's inline comment in `parse_toc` for the trade-off).
"""

import re
import zipfile
from io import BytesIO

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# A leading number, then an optional ".", ")", or ":", then whitespace, then the title text — the
# same pragmatic regex-heuristic approach as `formatting.upload`'s citation-style guess: good
# enough for the common "1. Introduction" / "2) Literature Review" / "3 Conclusion" shapes, not a
# real outline-format parser.
_NUMBERED_ENTRY_RE = re.compile(r"^\d+[.):]?\s+(.+)$")

# Best-effort cleanup for Word-generated TOC entries, which render as e.g.
# "Introduction .......... 5" (title, a dot-leader, then a page number) or, when the leader is a
# real tab character rather than literal dots, "Introduction\t5". Strips a trailing run of 1+
# dot/whitespace/tab characters followed by digits. This is a heuristic, not a guarantee: a title
# that legitimately ends in a number right after a short gap could be over-trimmed, but that's an
# acceptable trade-off for the common case this exists to handle.
_TRAILING_PAGE_NUMBER_RE = re.compile(r"[.\s]+\d+\s*$")

_HEADING_1_STYLE = "Heading 1"
_HEADING_2_STYLE = "Heading 2"

# Matches Word's built-in top-level TOC-field paragraph style, however it's spelled/spaced —
# `"TOC 1"`, `"TOC1"`, `"toc 1"`, etc. (Word's own naming varies by locale/version). Only level 1
# is matched, mirroring `_HEADING_1_STYLE`'s "top-level chapters only" scope: level 2+ entries
# (`"TOC 2"`, ...) are a document's subsections, not its chapters.
_TOC_FIELD_LEVEL_1_STYLE_RE = re.compile(r"^toc\s*1$", re.IGNORECASE)

# Detects a page-number suffix on its own, independent of a leading chapter number — a thesis's
# front/back matter entries (e.g. Russian "ВВЕДЕНИЕ"/"ЗАКЛЮЧЕНИЕ"/"СПИСОК ЛИТЕРАТУРЫ") are
# conventionally unnumbered even when the numbered chapters between them (e.g. "1 ...", "2 ...")
# are, so `_NUMBERED_ENTRY_RE` alone would silently drop them. Requires either a real tab
# character or a 2+-character dot-leader run before the trailing digits — deliberately stricter
# than `_TRAILING_PAGE_NUMBER_RE`'s cleanup regex, so a single space before an incidental trailing
# number in unrelated prose (e.g. "...took only 5") doesn't get mistaken for a TOC line.
_PAGE_NUMBER_SUFFIX_RE = re.compile(r"(\t+\d+|[.\s]{2,}\d+)\s*$")


def _has_page_number_suffix(text: str) -> bool:
    return _PAGE_NUMBER_SUFFIX_RE.search(text) is not None


# Matches a multi-level dotted subsection number ("1.1 ...", "2.3.1 ...") — always a subsection
# in this platform's chapter model (`Chapter`/subchapter, ADR-0014), never a top-level chapter,
# regardless of how the rest of the TOC is (un)structured. Captures the title text after the
# number so `parse_toc_with_subchapters` can use it directly, the same shape as
# `_NUMBERED_ENTRY_RE`'s capture group.
_DOTTED_SUBSECTION_RE = re.compile(r"^\d+(?:\.\d+)+[.):]?\s+(.+)$")

# Same heuristic as `projects.router._LEADING_CHAPTER_NUMBER_RE` (duplicated locally rather than
# imported, matching that module's own precedent of duplicating a small private helper across a
# module boundary rather than reaching into another module's private name): pulls a chapter's
# leading number out of its title regardless of the surrounding words. Used by
# `parse_document_sections_with_subchapters` to synthesize a numbered subsection title when the
# subsection's own paragraph text has no literal number of its own — see that function's inline
# comment for why that happens.
_CHAPTER_LEADING_NUMBER_RE = re.compile(r"^\D*(\d+)")

# Matches a lettered appendix sub-item ("Приложение А. ...", "Appendix A ...") — a sub-item of
# the single top-level "ПРИЛОЖЕНИЯ"/"Appendices" entry, not a chapter of its own.
_APPENDIX_SUBITEM_RE = re.compile(r"^(приложение|appendix)\s+[a-zа-яё0-9]+\b", re.IGNORECASE)

# Matches a per-chapter conclusion line ("Выводы по главе 1") — always a subsection of the
# chapter it summarizes, never a chapter of its own.
_CHAPTER_CONCLUSION_RE = re.compile(r"^выводы\s+по\s+глав", re.IGNORECASE)

# A TOC page's own heading ("ОГЛАВЛЕНИЕ"/"СОДЕРЖАНИЕ"/"Table of Contents"/"Contents") — never a
# chapter title itself, just the label of the page listing them.
_TOC_PAGE_TITLE_WORDS = frozenset({"оглавление", "содержание", "table of contents", "contents"})

# A real subsection *title* essentially never ends mid-sentence — unlike a subsection's *body*
# text, which (being ordinary prose) almost always does. Word documents frequently carry a
# "Heading 2"-style artifact onto the paragraph right after a real subsection heading (pasted
# text, "Format Painter", Ctrl+Shift+V not fully clearing formatting) — without this check,
# `parse_document_sections_with_subchapters` would misread that first body sentence as the start
# of a brand-new (bogus) subchapter, leaving the real subsection's own content empty (user
# report: "подглавы, которые написаны в документе, остаются пустыми"). Only applied to the
# `Heading 2`-style signal — a dotted-number match (`_DOTTED_SUBSECTION_RE`) is trusted
# regardless, since a body sentence starting with "1.1 " is implausible.
_SENTENCE_TERMINATOR_CHARS = ".!?…"


def _looks_like_a_heading(text: str) -> bool:
    return not text.rstrip("»\"'”").endswith(tuple(_SENTENCE_TERMINATOR_CHARS))


def _is_toc_subsection_or_page_title(text: str) -> bool:
    """`True` if `text` is a recognizable subsection/noise line that should never be treated as
    a top-level chapter title, even under `parse_toc`'s last-resort "every remaining line is a
    title" pass (see that pass's docstring for why it exists and why this filter matters)."""
    stripped = text.strip()
    if stripped.casefold() in _TOC_PAGE_TITLE_WORDS:
        return True
    return bool(
        _DOTTED_SUBSECTION_RE.match(stripped)
        or _APPENDIX_SUBITEM_RE.match(stripped)
        or _CHAPTER_CONCLUSION_RE.match(stripped)
    )


def _clean_toc_entry_title(text: str) -> str:
    """Strip a trailing tab-separated or dot-leader-separated page number from one TOC entry's
    raw paragraph text (see `_TRAILING_PAGE_NUMBER_RE`'s docstring), then strip whitespace.
    """
    if "\t" in text:
        text = text.split("\t", 1)[0]
    return _TRAILING_PAGE_NUMBER_RE.sub("", text).strip()


def _num_pr(paragraph_properties) -> object | None:
    return paragraph_properties.find(qn("w:numPr")) if paragraph_properties is not None else None


def _list_numbering_level(paragraph: Paragraph) -> int | None:
    """Returns `paragraph`'s list-numbering indent level (`0` = top level), or `None` if it isn't
    part of a numbered/bulleted list at all (no `w:numPr` anywhere in its own paragraph
    properties or its paragraph style's, walking the style-inheritance chain).

    Word's built-in list-style paragraph styles (e.g. "List Number") — which is what a real
    thesis TOC built by clicking a numbered-list button, rather than typing "1. " by hand, ends
    up using — put `w:numPr` on the *style definition* (`styles.xml`), not on each individual
    paragraph's own properties. A paragraph using such a style has no `w:numPr` of its own at
    all, so checking only the paragraph's own properties would silently miss this common case;
    `_TOC_FIELD_LEVEL_1_STYLE_RE`'s style-name check above only covers Word's specific `"TOC 1"`
    field style, not arbitrary list styles.

    A level-0 `w:ilvl` is often omitted entirely by Word (top level is the implicit default), so
    a present `w:numPr` with no `w:ilvl` child counts as level `0`, not `None`.
    """
    num_pr = _num_pr(paragraph._p.pPr)
    style = paragraph.style
    while num_pr is None and style is not None:
        style_element = getattr(style, "element", None)
        num_pr = _num_pr(style_element.find(qn("w:pPr")) if style_element is not None else None)
        style = getattr(style, "base_style", None)
    if num_pr is None:
        return None
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    try:
        return int(ilvl.get(qn("w:val")))
    except (TypeError, ValueError):
        return 0


class TocParseError(ValueError):
    """Raised when an uploaded file cannot be parsed into a table of contents.

    Callers (the TOC upload router) must translate this into a 4xx response rather than letting
    it surface as a 500 — this module fails closed rather than guessing a chapter list from a
    file it couldn't make sense of.
    """


def _top_level_entries(document) -> list[tuple[int, str]]:
    """Shared tiered-detection logic for `parse_toc`/`parse_toc_with_subchapters`: returns
    `(paragraph_index, title)` pairs for whichever tier (see module docstring) first yields any
    entries, preserving each entry's original paragraph position so
    `parse_toc_with_subchapters` can group subsection paragraphs under the nearest preceding
    entry regardless of which tier matched.
    """
    headings = [
        (index, paragraph.text.strip())
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.style is not None and paragraph.style.name == _HEADING_1_STYLE
    ]
    if headings:
        return headings

    toc_field_entries = [
        (index, _clean_toc_entry_title(paragraph.text))
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.style is not None
        and _TOC_FIELD_LEVEL_1_STYLE_RE.match(paragraph.style.name or "")
    ]
    toc_field_entries = [(index, title) for index, title in toc_field_entries if title]
    if toc_field_entries:
        return toc_field_entries

    list_numbered_entries = [
        (index, _clean_toc_entry_title(paragraph.text))
        for index, paragraph in enumerate(document.paragraphs)
        if _list_numbering_level(paragraph) == 0
    ]
    list_numbered_entries = [(index, title) for index, title in list_numbered_entries if title]
    if list_numbered_entries:
        return list_numbered_entries

    # Plain numbered lines (e.g. "1. Introduction") and unnumbered-but-page-numbered lines (e.g.
    # "ВВЕДЕНИЕ .......... 3") are scanned together, in one pass over document order, rather than
    # as two separate "first non-empty list wins" fallbacks: a real thesis TOC commonly mixes
    # both in the same document (numbered chapters alongside conventionally-unnumbered
    # introduction/conclusion/references entries), and if the numbered-only scan ran first and
    # found even one match, a strictly-sequential fallback would return early with just that
    # partial list, silently dropping every unnumbered entry instead of ever trying the
    # page-number-suffix check at all.
    entries = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        match = _NUMBERED_ENTRY_RE.match(text)
        if match is not None:
            title = _clean_toc_entry_title(match.group(1))
        elif _has_page_number_suffix(paragraph.text):
            # Requires the page-number-suffix signal specifically (see
            # `_PAGE_NUMBER_SUFFIX_RE`'s docstring) rather than treating every non-blank
            # paragraph as a title, so unrelated prose with no such structure still correctly
            # falls through to the last-resort tier below instead of being guessed at.
            title = _clean_toc_entry_title(paragraph.text)
        else:
            continue
        if title:
            entries.append((index, title))
    if entries:
        return entries

    # Absolute last resort: some real thesis TOCs (typed by hand, in "Normal" style throughout,
    # with no page numbers listed at all) carry no structural signal whatsoever beyond being a
    # short list of lines — every check above requires *some* marker (a style, a leading number,
    # a page-number suffix) and finds none. Since the caller specifically chose "upload a table
    # of contents" (a distinct, single-purpose upload from `parse_document_sections`'s "upload a
    # whole document", which never reaches this far — it requires `Heading 1` unconditionally),
    # that choice itself is the signal: trust every remaining non-blank line as a chapter title,
    # excluding only lines recognizable as a subsection/the page's own heading (see
    # `_is_toc_subsection_or_page_title`). This does mean a genuinely unrelated document dropped
    # into this upload gets misread as a one-line-per-chapter TOC rather than rejected — an
    # acceptable trade-off here specifically, since a wrong result just creates chapters the user
    # can immediately see and delete (`TASK-E11-3`), rather than silently corrupting anything.
    return [
        (index, paragraph.text.strip())
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() and not _is_toc_subsection_or_page_title(paragraph.text)
    ]


def _load_toc_document(content: bytes):
    try:
        return Document(BytesIO(content))
    except (PackageNotFoundError, KeyError, ValueError, zipfile.BadZipFile) as exc:
        raise TocParseError("Uploaded file is not a valid .docx document") from exc


def parse_toc(content: bytes) -> list[str]:
    """Parse `content` (raw `.docx` bytes) into an ordered list of chapter titles.

    Tried in order (see module docstring): `Heading 1` paragraphs, then Word TOC-field level-1
    paragraphs, then top-level Word-list-numbered paragraphs, then plain-numbered/page-numbered
    lines together, then (last resort) every remaining non-blank line. The first tier that
    yields any entries wins; a trailing dot-leader/tab/page-number artifact (e.g.
    `"Introduction .......... 5"` or `"Introduction\t5"`) is stripped from each entry as a
    best-effort cleanup (see `_TRAILING_PAGE_NUMBER_RE`).

    Raises `TocParseError` if `content` is not a valid `.docx` file, or if it has no non-blank
    paragraphs at all.
    """
    document = _load_toc_document(content)
    entries = _top_level_entries(document)
    if not entries:
        raise TocParseError(
            "Could not find any chapter titles in the .docx document — it appears to be empty."
        )
    return [title for _index, title in entries]


def parse_toc_with_subchapters(content: bytes) -> list[tuple[str, list[str]]]:
    """Parse `content` the same way as `parse_toc`, but also groups dotted-numbered subsection
    lines (e.g. `"3.1 ..."`, `"3.2 ..."` under a `"3 ..."`/`"ГЛАВА 3 ..."` top-level entry) as
    that entry's subchapters, instead of silently dropping them.

    Subsection detection (`_DOTTED_SUBSECTION_RE`) is independent of which top-level tier won —
    it never matches any of `parse_toc`'s own top-level patterns (a dotted `"3.1"` never starts a
    `Heading 1` paragraph, never gets a `"TOC 1"`-level style, is never list-numbered at level 0,
    and never matches `_NUMBERED_ENTRY_RE`'s single-number pattern) — so this runs as a second,
    orthogonal pass: every subsection paragraph is attached to the *nearest preceding* top-level
    entry in document order, which is always its actual parent chapter in a real TOC's document
    structure, regardless of numbering (a subsection's own leading number is not cross-checked
    against its parent's, since a TOC that skips/misnumbers a chapter would otherwise strand its
    subsections with no parent at all).

    Returns `[(title, [subchapter_title, ...]), ...]`, in document order; a chapter with no
    subsections gets an empty list. Raises `TocParseError` under the same conditions as
    `parse_toc`.
    """
    document = _load_toc_document(content)
    entries = _top_level_entries(document)
    if not entries:
        raise TocParseError(
            "Could not find any chapter titles in the .docx document — it appears to be empty."
        )

    chapters: list[tuple[str, list[str]]] = [(title, []) for _index, title in entries]
    entry_indices = [index for index, _title in entries]

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        if paragraph_index in entry_indices:
            continue
        stripped_text = paragraph.text.strip()
        if _DOTTED_SUBSECTION_RE.match(stripped_text) is None:
            continue
        # Keeps the leading number ("1.1 Онимы ...") in the subchapter title, unlike a top-level
        # chapter's number (stripped elsewhere) — a subsection's number is the reader's only way
        # to tell which parent chapter it belongs to and in what order, so displaying it is the
        # point (user request), not noise to clean up.
        subtitle = _clean_toc_entry_title(stripped_text)
        if not subtitle:
            continue
        # The chapter whose own index is the largest one still `<= paragraph_index` is this
        # subsection's nearest preceding top-level entry, i.e. its parent.
        parent_position = None
        for position, entry_index in enumerate(entry_indices):
            if entry_index <= paragraph_index:
                parent_position = position
            else:
                break
        if parent_position is not None:
            chapters[parent_position][1].append(subtitle)

    return chapters


def parse_document_sections(content: bytes) -> list[tuple[str, str]]:
    """Split a whole already-written `.docx` document into `(title, content)` sections, one per
    `Heading 1`-styled paragraph, so an entire pre-written thesis can be ingested as multiple
    chapters — each with its actual content, not just its title — in a single upload, instead of
    requiring the TOC (titles only, `parse_toc`) and then a separate per-chapter draft upload for
    each one.

    Every non-heading paragraph's text is appended to the section started by the most recent
    `Heading 1` paragraph above it; paragraphs before the first heading are dropped (there is no
    section yet to attach them to). Section content joins paragraphs with a blank line, mirroring
    a `.docx` document's paragraph-per-line structure. Blank paragraphs are skipped rather than
    preserved as empty lines.

    Raises `TocParseError` if `content` isn't a valid `.docx` file, or if no `Heading 1`
    paragraphs exist at all — unlike `parse_toc`, there is no numbered-list fallback here, since a
    numbered line has no notion of "everything until the next entry" to bound a section's content.
    """
    document = _load_toc_document(content)

    sections: list[tuple[str, list[str]]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        is_heading = paragraph.style is not None and paragraph.style.name == _HEADING_1_STYLE
        if is_heading:
            if text:
                sections.append((text, []))
            continue
        if not sections or not text:
            continue
        sections[-1][1].append(text)

    if not sections:
        raise TocParseError(
            "Could not find any Heading 1 paragraphs to split the document into chapters"
        )

    return [(title, "\n\n".join(paragraphs)) for title, paragraphs in sections]


def _normalize_heading_text(text: str) -> str:
    """Collapses internal whitespace runs (including a soft line break within a single
    paragraph — Word's Shift+Enter, which `python-docx` maps to a literal `"\\n"` mid-string per
    `Paragraph.text`'s own docstring) into single spaces, then strips. A long subsection heading
    wrapped onto two visual lines this way would otherwise defeat `_DOTTED_SUBSECTION_RE`'s
    unanchored-`.`-doesn't-match-newline regex match, and would look wrong as a title anyway with
    a raw newline embedded in it.
    """
    return " ".join(text.split())


def parse_document_sections_with_subchapters(
    content: bytes,
) -> list[tuple[str, str, list[tuple[str, str]]]]:
    """Split a whole already-written `.docx` document the same way as `parse_document_sections`,
    but also splits each chapter's own body into separate subchapters with their own content
    (e.g. `"3.1 ..."` under chapter `"3 ..."`, user request), instead of leaving them as
    undifferentiated body text under the parent chapter.

    A subsection boundary is recognized two ways, either sufficient on its own: a `Heading 2`-
    styled paragraph that also looks like a title rather than a sentence (see
    `_looks_like_a_heading` — guards against the common Word artifact where the paragraph right
    after a real subsection heading keeps its `Heading 2` style by mistake, which would otherwise
    misread that body text as a second, bogus subchapter and leave the real one empty), or — the
    same fallback `parse_toc_with_subchapters` uses for an unstyled TOC — a paragraph matching
    `_DOTTED_SUBSECTION_RE`'s dotted-number pattern regardless of style. A paragraph matching
    `_CHAPTER_CONCLUSION_RE` (e.g. `"Выводы по главе 1"`) is never itself treated as a
    subchapter — real theses commonly style it as `Heading 2` too, but it's a summary of the
    *chapter*, not a subsection of its own — and everything from it onward reverts to
    accumulating into the chapter's own content instead of whichever subsection preceded it.

    Paragraphs between a chapter's own `Heading 1` and its first subsection boundary (if any)
    become that chapter's own content; everything from a subsection boundary onward, up to the
    next boundary or the next `Heading 1`, becomes that subsection's content.

    A `Heading 2`-styled subsection whose own paragraph text has no literal number (Word's
    automatic multilevel-list numbering renders as "1.1" on screen but is never present in
    `paragraph.text` — the same phenomenon this module's own docstring describes for top-level
    chapters) gets one synthesized from its chapter's leading number plus a running count of its
    subsections so far, matching the numbering `parse_toc_with_subchapters` keeps in a TOC
    upload's subchapter titles — otherwise the two uploads' titles for the same real subsection
    could never match (`projects.router._match_existing_chapter`), permanently stranding the
    TOC's stub empty while the real content lands in an unmatched duplicate (user report).

    Returns `[(chapter_title, chapter_content, [(subchapter_title, subchapter_content), ...]),
    ...]`, in document order. `chapter_content` and each `subchapter_content` are empty strings
    when that section/subsection has no body paragraphs. Raises `TocParseError` under the same
    conditions as `parse_document_sections`.
    """
    document = _load_toc_document(content)

    chapters: list[dict] = []
    for paragraph in document.paragraphs:
        raw_text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else None

        if style_name == _HEADING_1_STYLE:
            if raw_text:
                chapters.append(
                    {
                        "title": _normalize_heading_text(raw_text),
                        "content": [],
                        "subchapters": [],
                        "in_conclusion": False,
                    }
                )
            continue
        if not chapters:
            continue
        current = chapters[-1]

        normalized = _normalize_heading_text(raw_text) if raw_text else ""
        dotted_match = _DOTTED_SUBSECTION_RE.match(normalized) if normalized else None
        is_subsection_boundary = bool(normalized) and (
            dotted_match is not None
            or (style_name == _HEADING_2_STYLE and _looks_like_a_heading(normalized))
        )

        if is_subsection_boundary and _CHAPTER_CONCLUSION_RE.match(normalized):
            current["in_conclusion"] = True
            continue

        if is_subsection_boundary and not current["in_conclusion"]:
            # Keeps the leading number for a dotted subsection ("1.1 Онимы ...") in its title —
            # see `parse_toc_with_subchapters`'s identical choice for why (user request).
            subtitle = _clean_toc_entry_title(normalized)
            if dotted_match is None:
                # Detected purely via the `Heading 2`-style signal, with no literal number
                # anywhere in the paragraph's own text — common when a thesis's body headings use
                # Word's automatic multilevel-list numbering, which (like the top-level-chapter
                # case this module's docstring already describes) renders as "1.1" on screen but
                # is never present in `paragraph.text`, since Word generates it from the list
                # definition rather than storing it as text. Synthesize the same "{chapter}.{n}"
                # numbering a TOC upload keeps in ITS subchapter titles (`parse_toc_with_
                # subchapters`), from this chapter's own leading number plus a running count of
                # its subsections so far — otherwise the two uploads' subchapter titles can never
                # match (`projects.router._match_existing_chapter`), leaving a TOC-created stub
                # empty forever while its real content lands in a second, duplicate subchapter
                # (user report).
                chapter_number_match = _CHAPTER_LEADING_NUMBER_RE.match(current["title"])
                if chapter_number_match is not None:
                    subtitle = f"{chapter_number_match.group(1)}.{len(current['subchapters']) + 1} {subtitle}"
            if subtitle:
                current["subchapters"].append({"title": subtitle, "content": []})
                continue

        if not raw_text:
            continue
        if current["in_conclusion"] or not current["subchapters"]:
            current["content"].append(raw_text)
        else:
            current["subchapters"][-1]["content"].append(raw_text)

    if not chapters:
        raise TocParseError(
            "Could not find any Heading 1 paragraphs to split the document into chapters"
        )

    return [
        (
            chapter["title"],
            "\n\n".join(chapter["content"]),
            [
                (subchapter["title"], "\n\n".join(subchapter["content"]))
                for subchapter in chapter["subchapters"]
            ],
        )
        for chapter in chapters
    ]
