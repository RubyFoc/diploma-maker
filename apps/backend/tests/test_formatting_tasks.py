"""Tests for `formatting.tasks.parse_toc_task` (ADR-0013, TASK-E17-2/TASK-E17-4).

`formatting.tasks.parse_toc_task` is a re-export of `toc.tasks.parse_toc_task` (see that module's
docstring for why the actual task is defined in `toc.tasks`). `parse_toc` is synchronous, so the
task calls it directly, with no `asyncio.run` involved — `.delay()` still needs to be called from
a plain sync test function to match every other task module's convention (ADR-0013 addendum
point 4), even though this particular task has no internal event loop to collide with.

The task's parameter is a base64-encoded string (TASK-E17-4), not raw bytes, so every `.delay()`
call here base64-encodes the `.docx` fixture bytes first, matching what
`projects.router.upload_toc_endpoint` does against a real broker.
"""

import base64
from io import BytesIO

import pytest
from docx import Document

from diploma_backend.formatting.tasks import parse_toc_task
from diploma_backend.toc.parser import TocParseError


def _docx_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_heading_docx() -> bytes:
    document = Document()
    document.add_paragraph("Introduction", style="Heading 1")
    document.add_paragraph("Literature Review", style="Heading 1")
    document.add_paragraph("Conclusion", style="Heading 1")
    return _docx_bytes(document)


def test_delay_runs_task_and_returns_titles() -> None:
    content_b64 = base64.b64encode(_build_heading_docx()).decode("ascii")
    async_result = parse_toc_task.delay(content_b64)

    assert async_result.get() == ["Introduction", "Literature Review", "Conclusion"]


def test_malformed_document_propagates_as_real_exception() -> None:
    content_b64 = base64.b64encode(b"not a real docx file").decode("ascii")
    with pytest.raises(TocParseError):
        parse_toc_task.delay(content_b64)
