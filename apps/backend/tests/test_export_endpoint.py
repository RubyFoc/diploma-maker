"""Tests for `GET /projects/{project_id}/export` (TASK-E06 closing the loop): the first HTTP
endpoint that actually reaches the docx export engine (`export/docx.py`) against a real project's
content.
"""

import asyncio
from io import BytesIO
from urllib.parse import unquote

from docx import Document
from fastapi.testclient import TestClient

from diploma_backend.db import get_database
from diploma_backend.formatting.models import (
    FontConfig,
    Headings,
    InstitutionConfig,
    MarginsMm,
    PageConfig,
)
from diploma_backend.formatting.service import create_institution_config
from diploma_backend.main import app

_DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _make_config(institution_id: str = "inst-1") -> InstitutionConfig:
    return InstitutionConfig(
        institution_id=institution_id,
        institution_name="Test University",
        source="seed",
        page=PageConfig(
            size="A4",
            orientation="portrait",
            margins_mm=MarginsMm(top=20, bottom=20, left=30, right=15),
        ),
        font=FontConfig(family="Times New Roman", size_pt=14, line_spacing=1.5),
        headings=Headings(),
        citation_style="APA",
        accuracy_weight=0.5,
        raw_sample_reference="sample-1",
    )


def _accept_chapter_content(client: TestClient, project_id: str, title: str, content: str) -> str:
    """Create a chapter under `project_id` and give it accepted `content`, bypassing the
    LLM-backed `/generate` endpoint by calling `versions.service` directly against the same
    in-memory fake DB the `client` fixture wires up.
    """
    chapter_id = client.post(
        f"/projects/{project_id}/chapters", json={"title": title}
    ).json()["id"]

    from diploma_backend.versions.service import accept_draft_version, create_draft_version

    db = app.dependency_overrides[get_database]()

    async def _create_and_accept() -> None:
        draft = await create_draft_version(db, chapter_id, content=content)
        await accept_draft_version(db, draft.id)

    asyncio.run(_create_and_accept())
    return chapter_id


def test_export_with_and_without_content_returns_valid_docx(client: TestClient) -> None:
    project_id = client.post("/projects", json={"title": "Thesis"}).json()["id"]
    _accept_chapter_content(client, project_id, "Introduction", "This chapter is done.")
    client.post(f"/projects/{project_id}/chapters", json={"title": "Conclusion"})

    response = client.get(f"/projects/{project_id}/export")

    assert response.status_code == 200
    assert response.headers["content-type"] == _DOCX_MEDIA_TYPE
    assert "attachment" in response.headers["content-disposition"]
    assert "Thesis" in response.headers["content-disposition"]

    document = Document(BytesIO(response.content))
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    all_text = "\n".join(p.text for p in document.paragraphs)

    assert headings == ["Introduction", "Conclusion"]
    assert "This chapter is done." in all_text
    assert "No accepted content yet." in all_text


def test_export_with_valid_institution_id_applies_styling(client: TestClient) -> None:
    project_id = client.post("/projects", json={"title": "Thesis"}).json()["id"]
    _accept_chapter_content(client, project_id, "Introduction", "Some content.")

    db = app.dependency_overrides[get_database]()
    asyncio.run(create_institution_config(db, _make_config("inst-1")))

    response = client.get(f"/projects/{project_id}/export", params={"institution_id": "inst-1"})

    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    normal_font = document.styles["Normal"].font
    assert normal_font.name == "Times New Roman"
    assert normal_font.size.pt == 14
    section = document.sections[0]
    assert round(section.top_margin.mm) == 20
    assert round(section.left_margin.mm) == 30


def test_export_with_unknown_institution_id_falls_back_without_404(client: TestClient) -> None:
    project_id = client.post("/projects", json={"title": "Thesis"}).json()["id"]
    _accept_chapter_content(client, project_id, "Introduction", "Some content.")

    response = client.get(
        f"/projects/{project_id}/export", params={"institution_id": "does-not-exist"}
    )

    assert response.status_code == 200
    document = Document(BytesIO(response.content))
    normal_font = document.styles["Normal"].font
    assert normal_font.name != "Times New Roman"


def test_export_404s_for_unknown_project(client: TestClient) -> None:
    response = client.get("/projects/does-not-exist/export")

    assert response.status_code == 404


def test_export_sanitizes_unsafe_characters_in_filename(client: TestClient) -> None:
    project_id = client.post(
        "/projects", json={"title": "My Thesis: Chapter 1/2?"}
    ).json()["id"]

    response = client.get(f"/projects/{project_id}/export")

    assert response.status_code == 200
    disposition = response.headers["content-disposition"]
    assert ":" not in disposition.split("filename=")[1]
    assert "/" not in disposition.split("filename=")[1]
    assert "?" not in disposition.split("filename=")[1]
    assert "My Thesis" in disposition


def test_export_preserves_cyrillic_title_in_filename(client: TestClient) -> None:
    """Regression test: an earlier version of `_sanitize_filename` allowlisted ASCII only, so a
    Cyrillic title (this platform's actual target audience per ADR-0001/sources.geo_filter) was
    turned into a string of underscores instead of a readable filename."""
    project_id = client.post(
        "/projects", json={"title": "Экспорт: Тест/Проверка?"}
    ).json()["id"]

    response = client.get(f"/projects/{project_id}/export")

    assert response.status_code == 200
    disposition = response.headers["content-disposition"]
    assert "filename*=UTF-8''" in disposition
    encoded = disposition.split("filename*=UTF-8''")[1]
    assert unquote(encoded) == "Экспорт_ Тест_Проверка_.docx"

    # Regression: the ASCII-only `filename=` fallback must never collapse to whitespace/nothing
    # for a fully non-ASCII title — it must fall back to the generic name instead.
    ascii_filename = disposition.split('filename="')[1].split('"')[0]
    assert any(char.isalnum() for char in ascii_filename)
