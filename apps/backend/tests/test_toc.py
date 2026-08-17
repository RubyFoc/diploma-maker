"""Tests for TASK-E10-2 (TOC upload + parser).

Builds tiny `.docx` fixtures on the fly with `python-docx`, matching
`test_formatting_upload.py`'s pattern, and exercises both `toc.parser.parse_toc` directly and the
`POST /projects/{project_id}/toc/upload` endpoint via `client` (see `conftest.py`).
"""

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from fastapi.testclient import TestClient

from diploma_backend.toc.parser import (
    TocParseError,
    parse_document_sections,
    parse_document_sections_with_subchapters,
    parse_toc,
    parse_toc_with_subchapters,
)


def _auth_headers(client: TestClient, email: str = "student@example.com") -> dict:
    """Register a user and return an `Authorization` header, since project endpoints require
    auth as of TASK-E11-1 (see `test_auth.py`'s `_register`)."""
    response = client.post("/auth/register", json={"email": email, "password": "hunter22"})
    assert response.status_code == 201
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def _docx_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_heading_docx() -> bytes:
    document = Document()
    document.add_paragraph("Introduction", style="Heading 1")
    document.add_paragraph("Some body text under the introduction.")
    document.add_paragraph("Literature Review", style="Heading 1")
    document.add_paragraph("Conclusion", style="Heading 1")
    return _docx_bytes(document)


def _build_heading_docx_with_subsections() -> bytes:
    """Mimics a real thesis document: `Heading 2` subsections within a `Heading 1` chapter,
    plus a "Выводы по главе N" (chapter conclusion) also styled `Heading 2` — which should stay
    part of the chapter's own content, not become its own subchapter."""
    document = Document()
    document.add_paragraph("ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ", style="Heading 1")
    document.add_paragraph("Вводный абзац главы.")
    document.add_paragraph("1.1 Первый подраздел", style="Heading 2")
    document.add_paragraph("Текст первого подраздела.")
    document.add_paragraph("1.2 Второй подраздел", style="Heading 2")
    document.add_paragraph("Текст второго подраздела.")
    document.add_paragraph("Выводы по главе 1", style="Heading 2")
    document.add_paragraph("Итоговый абзац главы.")
    document.add_paragraph("ГЛАВА 2 БЕЗ ПОДРАЗДЕЛОВ", style="Heading 1")
    document.add_paragraph("Просто текст без подразделов.")
    return _docx_bytes(document)


def _build_numbered_docx() -> bytes:
    document = Document()
    document.add_paragraph("Table of Contents")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("2. Literature Review")
    document.add_paragraph("3) Conclusion")
    return _docx_bytes(document)


def _build_numbered_with_page_numbers_docx() -> bytes:
    document = Document()
    document.add_paragraph("1. Introduction .......... 5")
    document.add_paragraph("2. Literature Review .......... 12")
    return _docx_bytes(document)


def _build_empty_docx() -> bytes:
    document = Document()
    document.add_paragraph("Just some unrelated prose.")
    document.add_paragraph("Nothing here looks like a TOC entry at all.")
    return _docx_bytes(document)


def _build_word_toc_field_docx() -> bytes:
    """Mimics a `.docx` containing just Word's auto-generated "Table of Contents" field, copied
    out on its own — the style Word actually gives each entry (`"TOC 1"`) rather than
    `"Heading 1"` or a manually-typed number, with a tab-separated page number the way Word's
    field renders it (not literal dot-leader text)."""
    document = Document()
    document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    document.styles.add_style("TOC 2", WD_STYLE_TYPE.PARAGRAPH)
    for title, page in [("Введение", 3), ("Обзор литературы", 5), ("Заключение", 20)]:
        paragraph = document.add_paragraph(title, style="TOC 1")
        paragraph.add_run(f"\t{page}")
    subsection = document.add_paragraph("Подраздел 1.1", style="TOC 2")
    subsection.add_run("\t6")
    return _docx_bytes(document)


def _build_list_numbered_docx() -> bytes:
    """Mimics a TOC where the numbering came from clicking Word's numbered-list button (style
    `"List Number"`) rather than typing "1. " by hand — the rendered number lives in the list
    definition (`w:numPr`), not in the paragraph's own text."""
    document = Document()
    document.add_paragraph("Введение", style="List Number")
    document.add_paragraph("Обзор литературы", style="List Number")
    document.add_paragraph("Заключение", style="List Number")
    return _docx_bytes(document)


def _build_mixed_numbered_and_unnumbered_docx() -> bytes:
    """Mimics a real Russian thesis TOC: numbered chapters plus conventionally-unnumbered
    front/back matter (introduction/conclusion/references), each with a dot-leader or
    tab-separated page number but no leading digit on the unnumbered entries."""
    document = Document()
    document.add_paragraph("ВВЕДЕНИЕ .......... 3")
    document.add_paragraph("1 Теоретические основы .......... 5")
    document.add_paragraph("2 Практическая часть .......... 15")
    document.add_paragraph("ЗАКЛЮЧЕНИЕ .......... 25")
    section = document.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    section.add_run("\t27")
    return _docx_bytes(document)


def _build_toc_with_dotted_subsections_docx() -> bytes:
    """No page numbers, no styling — the same unstructured shape as the real thesis TOC that
    motivated tier 6, but with dotted-numbered subsections ("1.1", "3.1"/"3.2") under two of the
    three chapters, and none under the middle one."""
    document = Document()
    document.add_paragraph("ВВЕДЕНИЕ")
    document.add_paragraph("ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    document.add_paragraph("1.1 Первый подраздел")
    document.add_paragraph("1.2 Второй подраздел")
    document.add_paragraph("ГЛАВА 2 БЕЗ ПОДРАЗДЕЛОВ")
    document.add_paragraph("ГЛАВА 3 ПРАКТИЧЕСКАЯ ЧАСТЬ")
    document.add_paragraph("3.1 Третий подраздел")
    document.add_paragraph("ЗАКЛЮЧЕНИЕ")
    return _docx_bytes(document)


def test_parse_toc_uses_heading_1_paragraphs() -> None:
    titles = parse_toc(_build_heading_docx())

    assert titles == ["Introduction", "Literature Review", "Conclusion"]


def test_parse_toc_falls_back_to_numbered_lines() -> None:
    titles = parse_toc(_build_numbered_docx())

    assert titles == ["Introduction", "Literature Review", "Conclusion"]


def test_parse_toc_strips_trailing_page_numbers() -> None:
    titles = parse_toc(_build_numbered_with_page_numbers_docx())

    assert titles == ["Introduction", "Literature Review"]


def test_parse_toc_includes_unnumbered_entries_alongside_numbered_ones() -> None:
    titles = parse_toc(_build_mixed_numbered_and_unnumbered_docx())

    assert titles == [
        "ВВЕДЕНИЕ",
        "Теоретические основы",
        "Практическая часть",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    ]


def test_parse_toc_invalid_file_raises() -> None:
    try:
        parse_toc(b"not a real docx file")
        raise AssertionError("expected TocParseError")
    except TocParseError:
        pass


def test_parse_toc_falls_back_to_every_line_when_nothing_else_matches() -> None:
    """Tier 6 (last resort, see module docstring): a document with no `Heading 1`/TOC-field/
    list-numbering/numbered-line/page-number-suffix signal at all still yields a chapter per
    non-blank line, trusting the caller's choice of "upload a table of contents" over raising."""
    titles = parse_toc(_build_empty_docx())

    assert titles == ["Just some unrelated prose.", "Nothing here looks like a TOC entry at all."]


def test_parse_toc_raises_for_a_document_with_no_non_blank_paragraphs() -> None:
    document = Document()
    document.add_paragraph("")
    document.add_paragraph("   ")

    try:
        parse_toc(_docx_bytes(document))
        raise AssertionError("expected TocParseError")
    except TocParseError:
        pass


def test_parse_toc_excludes_subsections_appendix_subitems_and_toc_page_title() -> None:
    document = Document()
    document.add_paragraph("ОГЛАВЛЕНИЕ")
    document.add_paragraph("ВВЕДЕНИЕ")
    document.add_paragraph("ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    document.add_paragraph("1.1 Подраздел первый")
    document.add_paragraph("1.2 Подраздел второй")
    document.add_paragraph("Выводы по главе 1")
    document.add_paragraph("ЗАКЛЮЧЕНИЕ")
    document.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    document.add_paragraph("ПРИЛОЖЕНИЯ")
    document.add_paragraph("Приложение А. Первое приложение")
    document.add_paragraph("Приложение Б. Второе приложение")

    titles = parse_toc(_docx_bytes(document))

    assert titles == [
        "ВВЕДЕНИЕ",
        "ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "ПРИЛОЖЕНИЯ",
    ]


def test_parse_toc_uses_word_toc_field_level_1_entries() -> None:
    titles = parse_toc(_build_word_toc_field_docx())

    assert titles == ["Введение", "Обзор литературы", "Заключение"]


def test_parse_toc_uses_top_level_list_numbered_paragraphs() -> None:
    titles = parse_toc(_build_list_numbered_docx())

    assert titles == ["Введение", "Обзор литературы", "Заключение"]


def test_parse_toc_with_subchapters_groups_dotted_sections_under_their_chapter() -> None:
    chapters = parse_toc_with_subchapters(_build_toc_with_dotted_subsections_docx())

    assert chapters == [
        ("ВВЕДЕНИЕ", []),
        ("ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ", ["1.1 Первый подраздел", "1.2 Второй подраздел"]),
        ("ГЛАВА 2 БЕЗ ПОДРАЗДЕЛОВ", []),
        ("ГЛАВА 3 ПРАКТИЧЕСКАЯ ЧАСТЬ", ["3.1 Третий подраздел"]),
        ("ЗАКЛЮЧЕНИЕ", []),
    ]


def test_parse_toc_with_subchapters_matches_parse_toc_when_nothing_is_nested() -> None:
    chapters = parse_toc_with_subchapters(_build_numbered_docx())

    assert [title for title, _subchapters in chapters] == parse_toc(_build_numbered_docx())
    assert all(subchapters == [] for _title, subchapters in chapters)


def test_upload_toc_creates_subchapters_for_dotted_sections(client: TestClient) -> None:
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    response = client.post(
        f"/projects/{project_id}/toc/upload",
        files={
            "file": (
                "toc.docx",
                _build_toc_with_dotted_subsections_docx(),
                "application/vnd.openxmlformats",
            )
        },
        headers=headers,
    )

    assert response.status_code == 201
    body = response.json()
    chapters_by_title = {chapter["title"]: chapter for chapter in body["chapters"]}
    assert set(chapters_by_title) == {
        "ВВЕДЕНИЕ",
        "ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ",
        "ГЛАВА 2 БЕЗ ПОДРАЗДЕЛОВ",
        "ГЛАВА 3 ПРАКТИЧЕСКАЯ ЧАСТЬ",
        "ЗАКЛЮЧЕНИЕ",
    }

    chapter_1_id = chapters_by_title["ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ"]["id"]
    subchapters_1 = client.get(
        f"/projects/{project_id}/chapters/{chapter_1_id}/subchapters", headers=headers
    ).json()
    assert [sub["title"] for sub in subchapters_1] == ["1.1 Первый подраздел", "1.2 Второй подраздел"]
    assert all(sub["parent_chapter_id"] == chapter_1_id for sub in subchapters_1)

    chapter_2_id = chapters_by_title["ГЛАВА 2 БЕЗ ПОДРАЗДЕЛОВ"]["id"]
    subchapters_2 = client.get(
        f"/projects/{project_id}/chapters/{chapter_2_id}/subchapters", headers=headers
    ).json()
    assert subchapters_2 == []

    chapter_3_id = chapters_by_title["ГЛАВА 3 ПРАКТИЧЕСКАЯ ЧАСТЬ"]["id"]
    subchapters_3 = client.get(
        f"/projects/{project_id}/chapters/{chapter_3_id}/subchapters", headers=headers
    ).json()
    assert [sub["title"] for sub in subchapters_3] == ["3.1 Третий подраздел"]


def test_upload_toc_creates_chapters_in_order(client: TestClient) -> None:
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    response = client.post(
        f"/projects/{project_id}/toc/upload",
        files={"file": ("toc.docx", _build_numbered_docx(), "application/vnd.openxmlformats")},
        headers=headers,
    )

    assert response.status_code == 201
    body = response.json()
    assert body["id"] == project_id
    titles = [chapter["title"] for chapter in body["chapters"]]
    assert titles == ["Introduction", "Literature Review", "Conclusion"]
    orders = [chapter["order"] for chapter in body["chapters"]]
    assert orders == sorted(orders)


def test_upload_toc_nonexistent_project_404s(client: TestClient) -> None:
    headers = _auth_headers(client)
    response = client.post(
        "/projects/does-not-exist/toc/upload",
        files={"file": ("toc.docx", _build_numbered_docx(), "application/vnd.openxmlformats")},
        headers=headers,
    )

    assert response.status_code == 404


def test_upload_toc_invalid_file_422s(client: TestClient) -> None:
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    response = client.post(
        f"/projects/{project_id}/toc/upload",
        files={"file": ("toc.docx", b"not a real docx file", "application/octet-stream")},
        headers=headers,
    )

    assert response.status_code == 422


# --- parse_document_sections / POST /{project_id}/document/upload (whole-document ingestion) ----


def test_parse_document_sections_splits_by_heading_1() -> None:
    sections = parse_document_sections(_build_heading_docx())

    assert sections == [
        ("Introduction", "Some body text under the introduction."),
        ("Literature Review", ""),
        ("Conclusion", ""),
    ]


def test_parse_document_sections_joins_multiple_paragraphs_per_section() -> None:
    document = Document()
    document.add_paragraph("Introduction", style="Heading 1")
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")

    sections = parse_document_sections(_docx_bytes(document))

    assert sections == [("Introduction", "First paragraph.\n\nSecond paragraph.")]


def test_parse_document_sections_drops_preamble_before_first_heading() -> None:
    document = Document()
    document.add_paragraph("Some cover-page text before any heading.")
    document.add_paragraph("Introduction", style="Heading 1")
    document.add_paragraph("Body text.")

    sections = parse_document_sections(_docx_bytes(document))

    assert sections == [("Introduction", "Body text.")]


def test_parse_document_sections_invalid_file_raises() -> None:
    try:
        parse_document_sections(b"not a real docx file")
        raise AssertionError("expected TocParseError")
    except TocParseError:
        pass


def test_parse_document_sections_no_headings_raises() -> None:
    try:
        parse_document_sections(_build_empty_docx())
        raise AssertionError("expected TocParseError")
    except TocParseError:
        pass


def test_upload_document_creates_chapters_with_content(client: TestClient) -> None:
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    response = client.post(
        f"/projects/{project_id}/document/upload",
        files={"file": ("thesis.docx", _build_heading_docx(), "application/vnd.openxmlformats")},
        headers=headers,
    )

    assert response.status_code == 201
    body = response.json()
    chapters = {chapter["title"]: chapter for chapter in body["chapters"]}
    assert set(chapters) == {"Introduction", "Literature Review", "Conclusion"}
    assert chapters["Introduction"]["pending_draft"]["content"] == (
        "Some body text under the introduction."
    )
    assert chapters["Literature Review"]["pending_draft"] is None


def test_upload_document_reuses_chapters_from_a_prior_toc_upload(client: TestClient) -> None:
    """A TOC upload followed by a whole-document upload for the same thesis should fill in the
    TOC's chapters, not create duplicates alongside them (user request)."""
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    toc_document = Document()
    toc_document.add_paragraph("introduction")  # deliberately different case than the heading
    toc_document.add_paragraph("Literature Review")
    toc_document.add_paragraph("Conclusion")
    toc_response = client.post(
        f"/projects/{project_id}/toc/upload",
        files={"file": ("toc.docx", _docx_bytes(toc_document), "application/vnd.openxmlformats")},
        headers=headers,
    )
    assert toc_response.status_code == 201
    toc_chapter_ids = {chapter["title"]: chapter["id"] for chapter in toc_response.json()["chapters"]}

    document_response = client.post(
        f"/projects/{project_id}/document/upload",
        files={"file": ("thesis.docx", _build_heading_docx(), "application/vnd.openxmlformats")},
        headers=headers,
    )

    assert document_response.status_code == 201
    body = document_response.json()
    chapters = {chapter["title"]: chapter for chapter in body["chapters"]}
    # Exactly 3 chapters, not 6 — the document upload reused each of the TOC's chapters rather
    # than creating a second copy of every one.
    assert len(body["chapters"]) == 3
    assert chapters["introduction"]["id"] == toc_chapter_ids["introduction"]
    assert chapters["Literature Review"]["id"] == toc_chapter_ids["Literature Review"]
    assert chapters["Conclusion"]["id"] == toc_chapter_ids["Conclusion"]
    assert chapters["introduction"]["pending_draft"]["content"] == (
        "Some body text under the introduction."
    )


def test_upload_document_reuses_toc_chapter_matched_only_by_leading_number(
    client: TestClient,
) -> None:
    """A real thesis's TOC entry and its document's actual `Heading 1` text frequently word the
    rest of the title differently even when they share the same chapter number — an exact-text
    match alone would miss this and create a duplicate (user report)."""
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    toc_document = Document()
    toc_document.add_paragraph("ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ИЗУЧЕНИЯ ВОПРОСА")
    toc_response = client.post(
        f"/projects/{project_id}/toc/upload",
        files={"file": ("toc.docx", _docx_bytes(toc_document), "application/vnd.openxmlformats")},
        headers=headers,
    )
    assert toc_response.status_code == 201
    toc_chapter_id = toc_response.json()["chapters"][0]["id"]

    document = Document()
    document.add_paragraph(
        "ГЛАВА 1 ТЕОРЕТИКО-МЕТОДОЛОГИЧЕСКИЕ ПРЕДПОСЫЛКИ ИССЛЕДОВАНИЯ", style="Heading 1"
    )
    document.add_paragraph("Содержание главы.")

    document_response = client.post(
        f"/projects/{project_id}/document/upload",
        files={"file": ("thesis.docx", _docx_bytes(document), "application/vnd.openxmlformats")},
        headers=headers,
    )

    assert document_response.status_code == 201
    body = document_response.json()
    assert len(body["chapters"]) == 1
    assert body["chapters"][0]["id"] == toc_chapter_id
    assert body["chapters"][0]["pending_draft"]["content"] == "Содержание главы."


def test_parse_document_sections_with_subchapters_splits_heading_2_and_excludes_conclusion() -> (
    None
):
    chapters = parse_document_sections_with_subchapters(_build_heading_docx_with_subsections())

    assert [title for title, _content, _subs in chapters] == [
        "ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ",
        "ГЛАВА 2 БЕЗ ПОДРАЗДЕЛОВ",
    ]
    chapter_1 = chapters[0]
    # The intro paragraph (before any subsection) and the "Выводы по главе 1" paragraph after it
    # (which reverts to the chapter's own content bucket, not the last subsection's) both land
    # in the chapter's own content.
    assert chapter_1[1] == "Вводный абзац главы.\n\nИтоговый абзац главы."
    assert [(title, content) for title, content in chapter_1[2]] == [
        ("1.1 Первый подраздел", "Текст первого подраздела."),
        ("1.2 Второй подраздел", "Текст второго подраздела."),
    ]
    # "Выводы по главе 1" is never its own subchapter, and its content doesn't leak into the
    # last subsection — it reverts to the chapter's own content bucket.
    subchapter_titles = [title for title, _content in chapter_1[2]]
    assert "Выводы по главе 1" not in subchapter_titles

    chapter_2 = chapters[1]
    assert chapter_2[1] == "Просто текст без подразделов."
    assert chapter_2[2] == []


def test_upload_document_creates_subchapters_from_heading_2(client: TestClient) -> None:
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    response = client.post(
        f"/projects/{project_id}/document/upload",
        files={
            "file": (
                "thesis.docx",
                _build_heading_docx_with_subsections(),
                "application/vnd.openxmlformats",
            )
        },
        headers=headers,
    )

    assert response.status_code == 201
    chapters_by_title = {chapter["title"]: chapter for chapter in response.json()["chapters"]}
    chapter_1 = chapters_by_title["ГЛАВА 1 ТЕОРЕТИЧЕСКИЕ ОСНОВЫ"]
    assert chapter_1["pending_draft"]["content"] == "Вводный абзац главы.\n\nИтоговый абзац главы."

    subchapters = client.get(
        f"/projects/{project_id}/chapters/{chapter_1['id']}/subchapters", headers=headers
    ).json()
    assert [sub["title"] for sub in subchapters] == ["1.1 Первый подраздел", "1.2 Второй подраздел"]

    chapter_2 = chapters_by_title["ГЛАВА 2 БЕЗ ПОДРАЗДЕЛОВ"]
    subchapters_2 = client.get(
        f"/projects/{project_id}/chapters/{chapter_2['id']}/subchapters", headers=headers
    ).json()
    assert subchapters_2 == []


def test_upload_document_nonexistent_project_404s(client: TestClient) -> None:
    headers = _auth_headers(client)
    response = client.post(
        "/projects/does-not-exist/document/upload",
        files={"file": ("thesis.docx", _build_heading_docx(), "application/vnd.openxmlformats")},
        headers=headers,
    )

    assert response.status_code == 404


def test_upload_document_no_headings_422s(client: TestClient) -> None:
    headers = _auth_headers(client)
    create_response = client.post("/projects", json={"title": "My Thesis"}, headers=headers)
    project_id = create_response.json()["id"]

    response = client.post(
        f"/projects/{project_id}/document/upload",
        files={"file": ("thesis.docx", _build_empty_docx(), "application/vnd.openxmlformats")},
        headers=headers,
    )

    assert response.status_code == 422
