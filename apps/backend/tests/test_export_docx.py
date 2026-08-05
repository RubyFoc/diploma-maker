"""Tests for TASK-E06-1 (Markdown → `.docx` mapping engine).

Builds Markdown input strings, converts them with `markdown_to_docx`/`markdown_to_docx_bytes`,
and reads the result back with `python-docx`'s `Document(BytesIO(...))` to assert on structure
and content, following the fixture style in `test_formatting_upload.py`.
"""

from io import BytesIO

from docx import Document

from diploma_backend.export.docx import (
    apply_institution_config,
    markdown_to_docx,
    markdown_to_docx_bytes,
)
from diploma_backend.formatting.models import (
    FontConfig,
    Headings,
    HeadingStyle,
    InstitutionConfig,
    MarginsMm,
    PageConfig,
)


def _make_config(*, size="A4", orientation="portrait", h1_extra=None) -> InstitutionConfig:
    return InstitutionConfig(
        institution_id="inst-1",
        institution_name="Test University",
        source="seed",
        page=PageConfig(
            size=size,
            orientation=orientation,
            margins_mm=MarginsMm(top=20, bottom=20, left=30, right=15),
        ),
        font=FontConfig(family="Times New Roman", size_pt=14, line_spacing=1.5),
        headings=Headings(h1=HeadingStyle(**(h1_extra or {}))),
        citation_style="APA",
        accuracy_weight=0.5,
        raw_sample_reference="sample-1",
    )


def test_headings_all_three_levels() -> None:
    document = markdown_to_docx("# Title\n\n## Section\n\n### Subsection\n")

    styles = [paragraph.style.name for paragraph in document.paragraphs]
    texts = [paragraph.text for paragraph in document.paragraphs]

    assert styles == ["Heading 1", "Heading 2", "Heading 3"]
    assert texts == ["Title", "Section", "Subsection"]


def test_plain_paragraph() -> None:
    document = markdown_to_docx("This is a plain paragraph of thesis text.")

    assert len(document.paragraphs) == 1
    paragraph = document.paragraphs[0]
    assert paragraph.text == "This is a plain paragraph of thesis text."
    assert paragraph.style.name in ("Normal", "Default Paragraph Font")


def test_multiline_paragraph_block_joins_with_space() -> None:
    document = markdown_to_docx("Line one\nLine two continues.\n\nSecond paragraph.")

    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "Line one Line two continues."
    assert document.paragraphs[1].text == "Second paragraph."


def test_bold_and_italic_inline_formatting() -> None:
    document = markdown_to_docx("This has **bold text** and *italic text* inside it.")

    paragraph = document.paragraphs[0]
    runs = paragraph.runs

    assert [run.text for run in runs] == [
        "This has ",
        "bold text",
        " and ",
        "italic text",
        " inside it.",
    ]
    assert runs[0].bold is not True
    assert runs[1].bold is True
    assert runs[1].italic is not True
    assert runs[2].bold is not True
    assert runs[3].italic is True
    assert runs[3].bold is not True
    assert runs[4].bold is not True


def test_unordered_list() -> None:
    document = markdown_to_docx("- First item\n* Second item\n- Third item")

    assert [p.style.name for p in document.paragraphs] == ["List Bullet"] * 3
    assert [p.text for p in document.paragraphs] == [
        "First item",
        "Second item",
        "Third item",
    ]


def test_ordered_list() -> None:
    document = markdown_to_docx("1. First step\n2. Second step\n3. Third step")

    assert [p.style.name for p in document.paragraphs] == ["List Number"] * 3
    assert [p.text for p in document.paragraphs] == [
        "First step",
        "Second step",
        "Third step",
    ]


def test_unsupported_construct_renders_as_plain_text_not_dropped() -> None:
    markdown_text = (
        "> This is a blockquote.\n\n| col1 | col2 |\n| --- | --- |\n| a | b |"
    )

    document = markdown_to_docx(markdown_text)
    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "This is a blockquote." in all_text
    assert "col1" in all_text
    assert "col2" in all_text
    # Neither construct should be turned into a heading or list style; it stays plain text.
    assert all(
        paragraph.style.name in ("Normal", "Default Paragraph Font")
        for paragraph in document.paragraphs
    )


def test_figure_placeholder_renders_as_own_italicized_paragraph() -> None:
    document = markdown_to_docx("[[figure: bar chart of quarterly revenue]]")

    assert len(document.paragraphs) == 1
    paragraph = document.paragraphs[0]
    assert paragraph.text == "[FIGURE PLACEHOLDER: bar chart of quarterly revenue]"
    assert len(paragraph.runs) == 1
    assert paragraph.runs[0].italic is True
    assert paragraph.style.name in ("Normal", "Default Paragraph Font")


def test_figure_placeholder_flushes_surrounding_paragraph_text() -> None:
    markdown_text = (
        "Some text before the figure.\n"
        "[[figure: diagram of system architecture]]\n"
        "Some text after the figure."
    )

    document = markdown_to_docx(markdown_text)

    assert len(document.paragraphs) == 3
    assert document.paragraphs[0].text == "Some text before the figure."
    assert document.paragraphs[1].text == "[FIGURE PLACEHOLDER: diagram of system architecture]"
    assert document.paragraphs[1].runs[0].italic is True
    assert document.paragraphs[2].text == "Some text after the figure."


def test_markdown_to_docx_bytes_round_trips_via_python_docx() -> None:
    docx_bytes = markdown_to_docx_bytes("# Heading\n\nSome **bold** text.")

    assert isinstance(docx_bytes, bytes)
    document = Document(BytesIO(docx_bytes))

    assert document.paragraphs[0].style.name == "Heading 1"
    assert document.paragraphs[0].text == "Heading"
    assert document.paragraphs[1].text == "Some bold text."


def test_apply_institution_config_a4_portrait_page_and_margins() -> None:
    document = markdown_to_docx("Some text.")
    apply_institution_config(document, _make_config(size="A4", orientation="portrait"))

    section = document.sections[0]
    assert round(section.page_width.mm, 1) == 210.0
    assert round(section.page_height.mm, 1) == 297.0
    assert round(section.top_margin.mm) == 20
    assert round(section.bottom_margin.mm) == 20
    assert round(section.left_margin.mm) == 30
    assert round(section.right_margin.mm) == 15


def test_apply_institution_config_a4_landscape_swaps_width_and_height() -> None:
    document = markdown_to_docx("Some text.")
    apply_institution_config(document, _make_config(size="A4", orientation="landscape"))

    section = document.sections[0]
    assert round(section.page_width.mm, 1) == 297.0
    assert round(section.page_height.mm, 1) == 210.0


def test_apply_institution_config_sets_normal_style_font() -> None:
    document = markdown_to_docx("Some text.")
    apply_institution_config(document, _make_config())

    normal_font = document.styles["Normal"].font
    assert normal_font.name == "Times New Roman"
    assert normal_font.size.pt == 14
    assert document.styles["Normal"].paragraph_format.line_spacing == 1.5


def test_apply_institution_config_heading_font_size_from_extra_fields() -> None:
    document = markdown_to_docx("# Title\n")
    config = _make_config(h1_extra={"font_size_pt": 18, "bold": True})
    apply_institution_config(document, config)

    heading_font = document.styles["Heading 1"].font
    assert heading_font.size.pt == 18
    assert heading_font.bold is True


def test_apply_institution_config_heading_ignores_unknown_extra_keys() -> None:
    document = markdown_to_docx("# Title\n")
    config = _make_config(h1_extra={"color": "red", "font_size_pt": 16})
    # Should not raise despite the unknown "color" key.
    apply_institution_config(document, config)

    assert document.styles["Heading 1"].font.size.pt == 16
